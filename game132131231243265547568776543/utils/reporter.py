from docx import Document
from datetime import datetime


def create_test_case_template():
    """
    Создаёт шаблон 'ТестКейс.docx' с метками {{RESULT_X}}.
    """
    doc = Document()
    doc.add_heading('Отчёт по тестированию системы', 0)
    doc.add_paragraph(f"Автоматически сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Действие'
    hdr_cells[1].text = 'Ожидаемый результат'
    hdr_cells[2].text = 'Результат'
    hdr_cells[3].text = 'Время'


    row = table.add_row().cells
    row[0].text = "Получить ФИО клиента"
    row[1].text = "ФИО в формате 'Фамилия& Имя& Отчество'"
    row[2].text = "{{RESULT_1}}"
    row[3].text = datetime.now().strftime("%H:%M:%S")


    row = table.add_row().cells
    row[0].text = "Проверка на запрещённые символы"
    row[1].text = "Нет недопустимых символов"
    row[2].text = "{{RESULT_2}}"
    row[3].text = datetime.now().strftime("%H:%M:%S")


    row = table.add_row().cells
    row[0].text = "Проверка на отсутствие цифр"
    row[1].text = "ФИО не содержит цифр"
    row[2].text = "{{RESULT_3}}"
    row[3].text = datetime.now().strftime("%H:%M:%S")

    doc.save('ТестКейс.docx')
    print("✅ Шаблон 'ТестКейс.docx' создан")


def _find_and_replace_in_paragraph(paragraph, placeholder, replacement):
    """
    Ищет и заменяет текст в параграфе, даже если он разбит на несколько run'ов.
    """
    if placeholder not in paragraph.text:
        return False

    full_text = paragraph.text
    if placeholder in full_text:
        new_text = full_text.replace(placeholder, replacement)

        for run in paragraph.runs:
            run.text = ""

        paragraph.runs[0].text = new_text
        return True
    return False


def update_result_in_docx(result_key, result_value):
    """
    Заменяет {{RESULT_X}} на значение во всём документе.
    """
    try:
        doc = Document('ТестКейс.docx')
        placeholder = f"{{{{{result_key}}}}}"
        found = False

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if _find_and_replace_in_paragraph(paragraph, placeholder, result_value):
                            found = True

        for paragraph in doc.paragraphs:
            if _find_and_replace_in_paragraph(paragraph, placeholder, result_value):
                found = True

        doc.save('ТестКейс.docx')

        if found:
            print(f"✅ Успешно заменено: {placeholder} → {result_value}")
        else:
            print(f"❌ Метка {placeholder} не найдена в документе")

        return found

    except Exception as e:
        print(f"❌ Ошибка при работе с документом: {e}")
        return False